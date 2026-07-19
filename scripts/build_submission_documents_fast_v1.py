#!/usr/bin/env python3
"""Build the fast, fixed-17-page DOCX submission draft.

The source Markdown is treated as content-only: its H1 and introductory build
note are intentionally excluded. Each H2 becomes one physical page and is
separated with an explicit Word page break.

Style basis: ``narrative_proposal`` with a named ``compact_17_page`` override:
Letter portrait, 1-inch margins, 10 pt body copy, single line spacing,
PingFang SC/Arial fonts, compact tables, and the MINISO red accent.
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "开题报告补充材料.md"
DEFAULT_OUTPUT = ROOT / "tmp" / "doc-fast-v1" / "名创优品_开题报告补充材料_fast_v1.docx"

PAGE_WIDTH_DXA = 9360
# LibreOffice's DOCX importer applies w:hAnsi more aggressively than Word on
# macOS, so PingFang SC is encoded into every font slot. It provides complete
# Simplified Chinese coverage and clean Latin glyphs in the rendered PDF.
BODY_FONT_CJK = "PingFang SC"
BODY_FONT_LATIN = "PingFang SC"
CODE_FONT = "PingFang SC"
MINISO_RED = "D71920"
INK = "242424"
MUTED = "666666"
LIGHT_FILL = "F7F3F3"
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "D8D8D8"


@dataclass(frozen=True)
class Section:
    title: str
    lines: tuple[str, ...]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return parser.parse_args()


def parse_sections(source: Path) -> list[Section]:
    lines = source.read_text(encoding="utf-8").splitlines()
    sections: list[Section] = []
    title: str | None = None
    body: list[str] = []

    for line in lines:
        if line.startswith("## "):
            if title is not None:
                sections.append(Section(title=title, lines=tuple(body)))
            title = line[3:].strip()
            body = []
        elif title is not None:
            body.append(line.rstrip())

    if title is not None:
        sections.append(Section(title=title, lines=tuple(body)))

    if len(sections) != 17:
        raise ValueError(f"expected 17 H2 sections, found {len(sections)}")
    for expected, section in enumerate(sections, start=1):
        match = re.match(r"第(\d+)页[｜|]", section.title)
        if not match or int(match.group(1)) != expected:
            raise ValueError(f"unexpected page title at position {expected}: {section.title}")
    return sections


def set_font(element, *, latin: str = BODY_FONT_LATIN, cjk: str = BODY_FONT_CJK) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), cjk)
    r_fonts.set(qn("w:cs"), latin)


def set_run_font(run, *, size: float = 10.0, code: bool = False) -> None:
    latin = CODE_FONT if code else BODY_FONT_LATIN
    cjk = BODY_FONT_CJK
    run.font.name = latin
    run.font.size = Pt(size)
    set_font(run._element, latin=latin, cjk=cjk)


def set_style_font(style, *, size: float, bold: bool = False, color: str = INK) -> None:
    style.font.name = BODY_FONT_LATIN
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    set_font(style.element, latin=BODY_FONT_LATIN, cjk=BODY_FONT_CJK)


def configure_styles(doc: DocumentObject) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=10.0, color=INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control = False

    heading = doc.styles["Heading 1"]
    set_style_font(heading, size=16.0, bold=True, color=MINISO_RED)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(5)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    label = doc.styles.add_style("Fast Label", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(label, size=10.0, bold=True, color=INK)
    label.paragraph_format.space_before = Pt(2)
    label.paragraph_format.space_after = Pt(1)
    label.paragraph_format.keep_with_next = True
    label.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    visual = doc.styles.add_style("Fast Visual Note", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(visual, size=9.2, color=MUTED)
    visual.paragraph_format.space_before = Pt(3)
    visual.paragraph_format.space_after = Pt(0)
    visual.paragraph_format.left_indent = Inches(0.08)
    visual.paragraph_format.right_indent = Inches(0.04)
    visual.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    visual.paragraph_format.keep_together = True

    list_style = doc.styles["List Bullet"]
    set_style_font(list_style, size=10.0, color=INK)
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(1)
    list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_document(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.45)

    doc.core_properties.title = "Trend2SKU 开题报告补充材料"
    doc.core_properties.subject = "名创优品 AI 驱动的产品开发智能决策引擎"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Trend2SKU, MINISO, Agent, 产品决策"
    doc.core_properties.comments = ""


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, value, end))
    set_run_font(run, size=8.0)


def add_footer(doc: DocumentObject) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    prefix = paragraph.add_run("TREND2SKU   |   ")
    set_run_font(prefix, size=8.0)
    prefix.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(paragraph, "PAGE")
    suffix = paragraph.add_run(" / 17")
    set_run_font(suffix, size=8.0)
    suffix.font.color.rgb = RGBColor.from_string(MUTED)


def set_paragraph_border(paragraph, *, color: str, side: str = "bottom", size: str = "14") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "3")
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def set_paragraph_fill(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def add_inline(paragraph, text: str, *, inherited_bold: bool = False, size: float = 10.0) -> None:
    """Add basic Markdown inline emphasis without flattening the text."""
    cursor = 0
    while cursor < len(text):
        bold_at = text.find("**", cursor)
        code_at = text.find("`", cursor)
        candidates = [(bold_at, "bold"), (code_at, "code")]
        candidates = [(position, kind) for position, kind in candidates if position >= 0]
        if not candidates:
            run = paragraph.add_run(text[cursor:])
            run.bold = inherited_bold
            set_run_font(run, size=size)
            break

        position, kind = min(candidates, key=lambda item: item[0])
        if position > cursor:
            run = paragraph.add_run(text[cursor:position])
            run.bold = inherited_bold
            set_run_font(run, size=size)

        if kind == "bold":
            end = text.find("**", position + 2)
            if end < 0:
                run = paragraph.add_run(text[position:])
                run.bold = inherited_bold
                set_run_font(run, size=size)
                break
            add_inline(
                paragraph,
                text[position + 2 : end],
                inherited_bold=True,
                size=size,
            )
            cursor = end + 2
        else:
            end = text.find("`", position + 1)
            if end < 0:
                run = paragraph.add_run(text[position:])
                run.bold = inherited_bold
                set_run_font(run, size=size)
                break
            run = paragraph.add_run(text[position + 1 : end])
            run.bold = inherited_bold
            run.font.color.rgb = RGBColor.from_string("8E1B1B")
            set_run_font(run, size=size, code=True)
            cursor = end + 1


def add_page_heading(doc: DocumentObject, title: str, page_number: int) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(2)
    kicker.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = kicker.add_run(f"TREND2SKU  ·  开题报告补充材料  ·  {page_number:02d}/17")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(MUTED)
    set_run_font(run, size=8.2)

    heading = doc.add_paragraph(style="Heading 1")
    add_inline(heading, title, inherited_bold=True, size=20.0 if page_number == 1 else 16.0)
    if page_number == 1:
        heading.paragraph_format.space_after = Pt(10)
    set_paragraph_border(heading, color=MINISO_RED, size="16")


def create_bullet_numbering(doc: DocumentObject) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.extend((start, num_fmt, lvl_text, lvl_jc))

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "260")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "20")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, indent, spacing))
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial Unicode MS")
    fonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend((ilvl, num_id_node))


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def visible_width(text: str) -> int:
    clean = re.sub(r"[*`]", "", text)
    return sum(2 if ord(char) > 127 else 1 for char in clean)


def allocate_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    column_count = len(rows[0])
    if column_count == 6:
        return [520, 2600, 1100, 820, 1260, 3060]
    if column_count == 4:
        return [3480, 1960, 1960, 1960]

    weights = []
    for column in range(column_count):
        maximum = max(visible_width(row[column]) for row in rows)
        weights.append(max(8, min(maximum, 34)))
    minimum = 980 if column_count == 3 else 1160
    remaining = PAGE_WIDTH_DXA - minimum * column_count
    total_weight = sum(weights)
    widths = [minimum + round(remaining * weight / total_weight) for weight in weights]
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def set_cell_margins(cell, *, top: int = 45, start: int = 80, bottom: int = 45, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), TABLE_BORDER)


def set_table_geometry(table, widths: Sequence[int]) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"table widths must total {PAGE_WIDTH_DXA}: {widths}")

    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_markdown_table(doc: DocumentObject, rows: Sequence[Sequence[str]]) -> None:
    widths = allocate_widths(rows)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.allow_autofit = False
    set_table_geometry(table, widths)
    set_table_borders(table)

    for row_index, (word_row, source_row) in enumerate(zip(table.rows, rows, strict=True)):
        for column_index, (cell, source_text) in enumerate(zip(word_row.cells, source_row, strict=True)):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, TABLE_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.widow_control = False
            plain = re.sub(r"[*`]", "", source_text)
            is_numeric = bool(re.fullmatch(r"[\[\]离线演示\s\d.%/–—-]+", plain))
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if row_index == 0 or is_numeric or column_index == 0 and len(plain) < 9
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline(paragraph, source_text, inherited_bold=row_index == 0, size=9.0)

    before = table._tbl.getprevious()
    if before is not None and before.tag == qn("w:p"):
        before.get_or_add_pPr()

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.line_spacing = Pt(2)


def iter_blocks(lines: Sequence[str]) -> Iterable[tuple[str, object]]:
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line:
            index += 1
            continue

        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            parsed = [split_table_row(item) for item in table_lines]
            if len(parsed) >= 2 and is_table_separator(table_lines[1]):
                parsed.pop(1)
            if not parsed or any(len(row) != len(parsed[0]) for row in parsed):
                raise ValueError(f"malformed Markdown table: {table_lines}")
            yield "table", parsed
            continue

        if line.startswith("- "):
            bullets: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("- "):
                bullets.append(lines[index].strip()[2:].strip())
                index += 1
            yield "bullets", bullets
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            following = lines[index].strip()
            if not following or following.startswith("|") or following.startswith("- "):
                break
            paragraph_lines.append(following)
            index += 1
        yield "paragraph", " ".join(paragraph_lines)


def add_paragraph_block(doc: DocumentObject, text: str) -> None:
    label_only = bool(re.fullmatch(r"\*\*[^*]+：\*\*", text))
    visual = text.startswith("**本页图表：**")
    if visual:
        text = text.replace("**本页图表：**", "**视觉呈现：**", 1)
        paragraph = doc.add_paragraph(style="Fast Visual Note")
        set_paragraph_fill(paragraph, LIGHT_FILL)
        set_paragraph_border(paragraph, color=MINISO_RED, side="left", size="18")
        add_inline(paragraph, text, size=9.2)
        return

    paragraph = doc.add_paragraph(style="Fast Label" if label_only else "Normal")
    add_inline(paragraph, text, size=10.0)


def add_section(doc: DocumentObject, section: Section, page_number: int, bullet_num_id: int) -> None:
    add_page_heading(doc, section.title, page_number)
    for block_type, payload in iter_blocks(section.lines):
        if block_type == "paragraph":
            add_paragraph_block(doc, str(payload))
        elif block_type == "bullets":
            for item in payload:
                paragraph = doc.add_paragraph(style="List Bullet")
                apply_numbering(paragraph, bullet_num_id)
                add_inline(paragraph, item, size=10.0)
        elif block_type == "table":
            add_markdown_table(doc, payload)
        else:
            raise AssertionError(f"unknown block type: {block_type}")


def build(source: Path, output: Path) -> Path:
    sections = parse_sections(source)
    doc = Document()
    configure_document(doc)
    configure_styles(doc)
    add_footer(doc)
    bullet_num_id = create_bullet_numbering(doc)

    for page_number, section in enumerate(sections, start=1):
        add_section(doc, section, page_number, bullet_num_id)
        if page_number != len(sections):
            break_paragraph = doc.add_paragraph()
            break_paragraph.paragraph_format.space_before = Pt(0)
            break_paragraph.paragraph_format.space_after = Pt(0)
            break_paragraph.add_run().add_break(WD_BREAK.PAGE)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def main() -> None:
    args = parse_args()
    output = build(args.source.resolve(), args.output.resolve())
    print(output)


if __name__ == "__main__":
    main()
